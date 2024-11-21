import os
import re

import config
from langchain.prompts import PromptTemplate
from langchain_openai import ChatOpenAI
from langchain_mistralai import ChatMistralAI
from langchain_google_genai import ChatGoogleGenerativeAI
#from langchain_community.chat_models.gigachat import GigaChat
from langchain_gigachat import GigaChat
#from yandex_chain import YandexLLM
from langchain_community.llms import YandexGPT
from langchain_core.prompts import ChatPromptTemplate, StringPromptTemplate, SystemMessagePromptTemplate, HumanMessagePromptTemplate
from langchain_community.document_loaders import (
    TextLoader,
    PDFMinerLoader,
    UnstructuredWordDocumentLoader,
    UnstructuredExcelLoader,
    UnstructuredPowerPointLoader,
    UnstructuredURLLoader
)
from langchain.document_loaders.base import BaseLoader
from docx import Document
from langchain.schema import Document as LangChainDocument

from abc import ABC, abstractmethod
from typing import List, Any, Optional, Dict, Tuple

import logging

from utils import transform_text_to_list, group_paragraphs
from prompts import (
    legal_prompt_agreement_type,
    
    legal_short_prompt,
    legal_warning
    )

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

os.environ["LANGCHAIN_TRACING_V2"] = "true"

def extract_text_with_correct_numbering(file_path):
    doc = Document(file_path)
    full_text = []
    counters = {}  # Dictionary to hold counts for each numId and ilvl
    previous_numId = None

    for para in doc.paragraphs:
        numbering = ""
        is_list = False

        # Check if the paragraph has numbering properties
        if para.style.name.startswith('List') or para._p.pPr and para._p.pPr.numPr:
            is_list = True
            try:
                num_pr = para._p.pPr.numPr
                num_id = num_pr.numId.val
                ilvl = num_pr.ilvl.val
            except AttributeError:
                num_id = 0
                ilvl = 0

            # Initialize counters for a new numId
            if num_id != previous_numId:
                previous_numId = num_id
                if num_id not in counters:
                    counters[num_id] = {}
                # Reset all deeper levels when a new list starts
                for lvl in list(counters[num_id].keys()):
                    if lvl >= ilvl:
                        counters[num_id][lvl] = 0

            # Initialize or increment the counter for the current level
            if ilvl not in counters[num_id]:
                counters[num_id][ilvl] = 1
            else:
                counters[num_id][ilvl] += 1

            # Reset counters for deeper levels
            levels_to_reset = [lvl for lvl in counters[num_id] if lvl > ilvl]
            for lvl in levels_to_reset:
                counters[num_id][lvl] = 0

            # Build the numbering string based on current counters
            numbering_parts = [str(counters[num_id][lvl]) for lvl in sorted(counters[num_id].keys()) if lvl <= ilvl and counters[num_id][lvl] > 0]
            numbering = '.'.join(numbering_parts) + '.'

            # Append the numbered text
            full_text.append(f"{numbering} {para.text}")
        else:
            # Reset previous_numId when encountering a non-list paragraph
            previous_numId = None
            full_text.append(para.text)

    combined_text = "\n".join(full_text)
    return combined_text

class PythonDocxLoader:
    def __init__(self, file_path: str):
        self.file_path = file_path

    def load(self):
        text = extract_text_with_correct_numbering(self.file_path)
        return [LangChainDocument(page_content=text)]

def get_loader(file_path):
    extension = os.path.splitext(file_path)[1].lower()
    if extension == ".txt":
        return TextLoader(file_path, encoding='utf-8')
    elif extension == ".pdf":
        return PDFMinerLoader(file_path)
    elif extension in [".docx", ".doc"]:
        return UnstructuredWordDocumentLoader(file_path)
        #return PythonDocxLoader(file_path)
    elif extension in [".xlsx", ".xls"]:
        return UnstructuredExcelLoader(file_path)
    elif extension in [".pptx", ".ppt"]:
        return UnstructuredPowerPointLoader(file_path)
    elif extension in [".url", ".html", ".htm"]:
        return UnstructuredURLLoader(file_path)
    else:
        raise ValueError(f"Unsupported file extension: {extension}")
    

from typing import Optional
from pydantic import BaseModel, Field, root_validator, model_validator
from enum import Enum

# Initialize the assistant with the system prompt
llm = ChatOpenAI(openai_api_key=config.OPENAI_API_KEY, model='gpt-4o-mini', temperature=0.1)

class Party(BaseModel):
    party_legal_name: Optional[str] = Field(..., description="Наименование стороны Договора")
    party_role: Optional[str] = Field(..., description="Наименование стороны в Договоре (например: именуемый в дальнейшем Заказчик)")

class Classification(BaseModel):
    document_type: str = Field(..., enum=["ДОГОВОР АРЕНДЫ", "ДОГОВОР ПОСТАВКИ", "ДОГОВОР ПОДРЯДА", "ДОГОВОР ОКАЗАНИЯ УСЛУГ", "АГЕНТСКИЙ ДОГОВОР", "ДРУГОЕ"])
    parties: Optional[List[Party]] = Field(..., description="Стороны Договора")

    @model_validator(mode='after')
    def set_default_roles_and_names(cls, model):
        if model.parties:
            for index, party in enumerate(model.parties):
                # Set party_role to "Party_<index>" if it's None or empty
                if not party.party_role:
                    party.party_role = f"Party_{index}"
                # Set party_legal_name to party_role if it's None or empty
                if not party.party_legal_name:
                    party.party_legal_name = party.party_role
        return model

tagging_prompt = ChatPromptTemplate.from_template(legal_prompt_agreement_type)
tagging_chain = llm.with_structured_output(Classification)
classify_chain = tagging_prompt | tagging_chain

for file_path in os.listdir('./tests/'):
    try:
        loader = get_loader(f'./tests/{file_path}')
        documents = loader.load()
        full_text = "\n".join([doc.page_content for doc in documents])
        full_text = ''.join(char for char in full_text if char.isprintable())
        pattern = r'(?m)(?=^\d+(?:\.\d+)*\.\s)'
        #paragraphs  = re.split(pattern, full_text)
        paragraphs = transform_text_to_list(full_text)
        if len(paragraphs) == 0:
            paragraphs = transform_text_to_list(full_text, add_lines=True)

        # Определяем тип документа
        zero_chunk = paragraphs[0]['text'] if paragraphs else full_text.strip()[:256]
        contract_metadata = classify_chain.invoke({"contract": zero_chunk})

        if contract_metadata.document_type != 'ДРУГОЕ':
            warnings = legal_warning[contract_metadata.document_type]
        else:
            raise Exception('Еще не реализовано.')
        
        if contract_metadata.parties is None:
            contract_metadata.parties = [Party(party_legal_name="Сторона 1", party_role="Сторона 1"), Party(party_legal_name="Сторона 2", party_role="Сторона 2")]
        for party in contract_metadata.parties:
            print(party.party_legal_name + " ИМЕНУЕМЫЙ " + party.party_role)

        legal_comments = []
        
        paras = group_paragraphs(paragraphs, 1)

        for paragraph in paras[1:]:
            human_prompt_template = "Верни комментарии для параграфа {paragraph} для стороны Договора {party}."
            system_message = SystemMessagePromptTemplate.from_template(legal_short_prompt)
            human_message = HumanMessagePromptTemplate.from_template(human_prompt_template)
            chat_prompt = ChatPromptTemplate.from_messages([system_message, human_message])
            chain = chat_prompt | llm
            comments = chain.invoke({"legal_warnings": warnings, "contract": full_text, "paragraph": paragraph['text'], "party": contract_metadata.parties[0]})    
            legal_comments.append(f'Комментарий к параграфу {paragraph["text"]}:\n{comments.content}\n\n') #legal_comment = result.content

        with open(f'./results/{file_path}.txt', 'w', encoding='utf-8') as f:
            f.write("\n\n".join(legal_comments))
        print(f"{file_path} processed.")
    except Exception as e:
        print(f"Error processing {file_path}: {str(e)}")

